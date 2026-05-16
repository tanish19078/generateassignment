import json
import re
import io
import time
import os
import gzip
import traceback
import urllib.error
import urllib.request
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from dotenv import load_dotenv

load_dotenv()
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


app = Flask(__name__, static_folder='public', static_url_path='')
CORS(app)


MAX_COMPRESSED_DOWNLOAD_BYTES = 3_500_000
MAX_DECOMPRESSED_DOWNLOAD_BYTES = 50_000_000


class DownloadRequestError(ValueError):
    pass


def read_download_payload():
    encoding = request.headers.get('X-Content-Encoding', '').strip().lower()
    body = request.get_data(cache=False)

    if encoding == 'gzip':
        if len(body) > MAX_COMPRESSED_DOWNLOAD_BYTES:
            raise DownloadRequestError('Compressed export payload is too large for this deployment.')

        try:
            body = gzip.decompress(body)
        except (OSError, EOFError) as err:
            raise DownloadRequestError('Compressed export payload could not be decoded.') from err
    elif encoding not in ('', 'identity'):
        raise DownloadRequestError(f'Unsupported export payload encoding: {encoding}')

    if not body:
        return None

    if len(body) > MAX_DECOMPRESSED_DOWNLOAD_BYTES:
        raise DownloadRequestError('Export payload is too large after decompression.')

    try:
        return json.loads(body.decode('utf-8'))
    except (UnicodeDecodeError, json.JSONDecodeError) as err:
        raise DownloadRequestError('Export payload is not valid JSON.') from err


LLM_PROVIDERS = {
    'groq': {
        'label': 'Groq',
        'base_url': 'https://api.groq.com/openai/v1',
        'env_vars': ['GROQ_API_KEY'],
        'api_format': 'openai',
        'requires_key': True,
    },
    'cerebras': {
        'label': 'Cerebras',
        'base_url': 'https://api.cerebras.ai/v1',
        'env_vars': ['CEREBRAS_API_KEY'],
        'api_format': 'openai',
        'requires_key': True,
    },
    'freemodel_openai': {
        'label': 'FreeModel OpenAI',
        'base_url': 'https://api.freemodel.dev/v1',
        'env_vars': ['FREEMODEL_OPENAI_API_KEY', 'FREEMODEL_API_KEY'],
        'api_format': 'openai',
        'requires_key': True,
    },
    'freemodel_anthropic': {
        'label': 'FreeModel Claude',
        'base_url': 'https://cc.freemodel.dev/v1',
        'env_vars': ['FREEMODEL_ANTHROPIC_API_KEY', 'FREEMODEL_API_KEY'],
        'api_format': 'anthropic',
        'anthropic_version': '2023-06-01',
        'fallback_model': 'claude-sonnet-4-6',
        'max_tokens': 4096,
        'requires_key': True,
    },
}


def get_env_values(env_names):
    values = []

    for env_name in env_names:
        value = os.getenv(env_name)
        if value:
            values.append(value)

    env_path = os.path.join(os.path.dirname(__file__), '.env')
    if os.path.exists(env_path):
        try:
            with open(env_path, 'r', encoding='utf-8') as env_file:
                for line in env_file:
                    stripped = line.strip()
                    if not stripped or stripped.startswith('#') or '=' not in stripped:
                        continue

                    key, value = stripped.split('=', 1)
                    if key.strip() in env_names and value.strip():
                        values.append(value.strip().strip('"').strip("'"))
        except OSError:
            pass

    deduped = []
    for value in values:
        if value not in deduped:
            deduped.append(value)

    return deduped


def get_provider_keys(provider_config, submitted_key):
    env_keys = get_env_values(provider_config.get('env_vars', []))
    if submitted_key:
        return [submitted_key] + [key for key in env_keys if key != submitted_key]

    return env_keys


def parse_openai_chat_response(provider_config, body):
    parsed = json.loads(body)
    try:
        return parsed['choices'][0]['message']['content']
    except (KeyError, IndexError, TypeError) as err:
        raise ValueError(f"Malformed {provider_config['label']} response: {body[:500]}") from err


def parse_anthropic_message_response(provider_config, body):
    parsed = json.loads(body)
    content = parsed.get('content')

    if isinstance(content, str):
        return content

    if isinstance(content, list):
        text_blocks = []
        for block in content:
            if isinstance(block, str):
                text_blocks.append(block)
            elif isinstance(block, dict) and block.get('text'):
                text_blocks.append(block['text'])

        if text_blocks:
            return '\n'.join(text_blocks)

    try:
        return parsed['choices'][0]['message']['content']
    except (KeyError, IndexError, TypeError) as err:
        raise ValueError(f"Malformed {provider_config['label']} response: {body[:500]}") from err


def build_anthropic_payload(provider_config, model, messages):
    system_parts = []
    anthropic_messages = []

    for message in messages:
        role = message.get('role', 'user')
        content = message.get('content', '')

        if role == 'system':
            system_parts.append(content)
            continue

        if role not in ('user', 'assistant'):
            role = 'user'

        anthropic_messages.append({
            'role': role,
            'content': content,
        })

    payload = {
        'model': model,
        'max_tokens': provider_config.get('max_tokens', 4096),
        'messages': anthropic_messages or [{'role': 'user', 'content': ''}],
    }

    if system_parts:
        payload['system'] = '\n\n'.join(system_parts)

    return payload


def create_chat_completion(provider_key, api_key, model, messages):
    provider_config = LLM_PROVIDERS.get(provider_key)
    if not provider_config:
        raise ValueError(f"Unsupported LLM provider: {provider_key}")

    if provider_config.get('requires_key') and not api_key:
        env_hint = ' or '.join(provider_config.get('env_vars', []))
        raise ValueError(f"{provider_config['label']} API key not found. Enter a key or set {env_hint}.")

    api_format = provider_config.get('api_format', 'openai')
    if api_format == 'anthropic':
        url = provider_config['base_url'].rstrip('/') + '/messages'
        payload = build_anthropic_payload(provider_config, model, messages)
    else:
        url = provider_config['base_url'].rstrip('/') + '/chat/completions'
        payload = {
            'model': model,
            'messages': messages,
        }

    headers = {
        'Accept': 'application/json',
        'Content-Type': 'application/json',
        'User-Agent': 'PractiGen/5.2',
    }
    if api_key and api_format == 'anthropic':
        headers['x-api-key'] = api_key
        headers['anthropic-version'] = provider_config.get('anthropic_version', '2023-06-01')
    elif api_key:
        headers['Authorization'] = f'Bearer {api_key}'

    req = urllib.request.Request(url, data=json.dumps(payload).encode('utf-8'), headers=headers, method='POST')

    try:
        with urllib.request.urlopen(req, timeout=120) as response:
            body = response.read().decode('utf-8')
    except urllib.error.HTTPError as err:
        body = err.read().decode('utf-8', errors='replace')
        raise RuntimeError(f"{provider_config['label']} API error {err.code}: {body}") from err
    except urllib.error.URLError as err:
        raise RuntimeError(f"{provider_config['label']} API connection error: {err.reason}") from err

    if api_format == 'anthropic':
        return parse_anthropic_message_response(provider_config, body)

    return parse_openai_chat_response(provider_config, body)

@app.route('/')
def serve_index():
    return send_from_directory('public', 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory('public', path)


# ==================== API: Parse Aims ====================
@app.route('/api/parse', methods=['POST'])
def api_parse():
    try:
        data = request.get_json()
        text = data.get('text', '')
        separator = data.get('separator', '---')

        pattern = r'\n\s*' + re.escape(separator) + r'+\s*\n'
        aim_blocks = re.split(pattern, text)
        aims = [b.strip() for b in aim_blocks if b.strip()]

        return jsonify({'aims': aims})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==================== STEP PARSING (OS MODE) ====================
def parse_steps(procedure_text):
    """Parse the procedure text into individual steps with explanation, command, and output.
    Handles both 'Step N:' format and 'N.' numbered list format from the LLM."""
    steps = []

    has_step_format = bool(re.search(r'Step\s+\d+\s*[:.]\s', procedure_text, re.IGNORECASE))

    if has_step_format:
        split_pat = r'(?=Step\s+\d+\s*[:.]\s)'
        header_pat = r'Step\s+(\d+)\s*[:.]\s*(.*?)(?:\n)'
    else:
        split_pat = r'(?:^|\n)(?=\d+\.\s)'
        header_pat = r'(\d+)\.\s*(.*?)(?:\n)'

    step_blocks = re.split(split_pat, procedure_text, flags=re.IGNORECASE)
    step_blocks = [b.strip() for b in step_blocks if b.strip()]

    print(f"DEBUG parse_steps: {len(step_blocks)} blocks (format: {'Step N' if has_step_format else 'N.'})")

    for idx, block in enumerate(step_blocks):
        header_match = re.match(header_pat, block, re.IGNORECASE)
        if not header_match:
            print(f"DEBUG: Block {idx} no header: {block[:80]}...")
            continue

        step_num = int(header_match.group(1))
        explanation = header_match.group(2).strip()
        explanation = re.sub(r'\*\*([^*]*)\*\*', r'\1', explanation)
        explanation = explanation.rstrip(':').strip()

        rest = block[header_match.end():]

        output_part = ''
        command_part = rest.strip()

        for pattern in [
            r'\n\s*\*{0,2}Output\*{0,2}\s*:\s*\n',
            r'\n\s*\*{0,2}Output\*{0,2}\s*:\s*',
            r'\*{0,2}Output\*{0,2}\s*:\s*\n',
        ]:
            parts = re.split(pattern, rest, maxsplit=1, flags=re.IGNORECASE)
            if len(parts) == 2:
                command_part = parts[0].strip()
                output_part = parts[1].strip()
                break
        else:
            print(f"DEBUG: Step {step_num} no Output: delimiter")

        command_part = re.sub(r'^\$\s*', '', command_part, flags=re.MULTILINE)
        command_part = re.sub(r'\*\*([^*]*)\*\*', r'\1', command_part)

        print(f"DEBUG: Step {step_num} cmd:{len(command_part)} out:{len(output_part)}")

        steps.append({
            'num': step_num,
            'explanation': explanation,
            'command': command_part,
            'output': output_part,
        })

    if not steps:
        print(f"DEBUG: No steps! Text: {procedure_text[:200]}...")
        steps = [{'num': 1, 'explanation': 'Execute the procedure', 'command': procedure_text, 'output': ''}]

    return steps


# ==================== API: Generate Content ====================
@app.route('/api/generate', methods=['POST'])
def api_generate():
    try:
        data = request.get_json()
        aim = data.get('aim', '')
        api_key = data.get('api_key', '')
        provider = data.get('provider', 'groq')
        model = data.get('model', 'llama-3.3-70b-versatile')
        mode = data.get('mode', 'general')
        target_language = data.get('code_language', '').strip()
        if mode in ('general', 'language') and not target_language:
            raise ValueError("Code language is required for General Coding mode.")
        
        terminal_user = data.get('terminal_user', 'student')
        if not terminal_user.strip():
            terminal_user = 'student'
            
        terminal_host = data.get('terminal_host', 'kali')
        if not terminal_host.strip():
            terminal_host = 'kali'

        selected_provider_config = LLM_PROVIDERS.get(provider)
        if not selected_provider_config:
            raise ValueError(f"Unsupported LLM provider: {provider}")
        selected_api_keys = get_provider_keys(selected_provider_config, api_key)

        if mode == 'os':
            prompt = f"""You are a professional Linux systems instructor preparing a practical lab file for a university Operating Systems course. Your instructor has assigned this aim:

"{aim}"

Work through the aim methodically, step by step, like an experienced professional would demonstrate in a real lab. Before each command or code block, write a clear explanation (1-2 lines) of what it does and why. Then write the command. Then show what the terminal actually displayed.

If the aim asks you to explore a command with its options, use it normally first, then show a few useful options — just like you'd actually try them in a real lab session. Don't robotically list every flag. Use your judgement.

If the aim asks for a C program (system calls, algorithms, etc.), write a clean, complete program. Compile and run it.

If the aim has multiple parts or multiple commands, work through each one properly.

Keep it natural. No filler. No padding. Just a real, useful practical.

Respond in this exact format:

[CONCEPT]
4-5 lines explaining the core OS concepts behind this practical in academic language. Keep it concise and focused. If there is a specific term or technique that needs extra explanation (e.g., what a system call is, what a process control block does), add ONE short follow-up paragraph for it — but only if truly needed. Do not over-explain obvious things.

[PROCEDURE]
Plain text, no markdown fences. Write it as numbered steps. Each step MUST include its own output immediately after the command.

Format each step EXACTLY like this:

Step 1: <what you're doing and why>
$ <command or code>
Output:
{terminal_user}@{terminal_host}:~$ <the command typed>
<realistic terminal output for this specific command>

Step 2: <what you're doing and why>
$ <command or code>
Output:
{terminal_user}@{terminal_host}:~$ <the command typed>
<realistic terminal output for this specific command>

...and so on. Only what the aim needs.

For C programs, write the full source code in one step (no $ prefix for the code itself), then compile and run as separate steps with $ prefix and their own outputs.

CRITICAL — Terminal output rules:
- The Output section for each step MUST start with the prompt and the command being typed on the FIRST line, then show the result below it. Do NOT add an extra prompt line at the end.
- Use this prompt: {terminal_user}@{terminal_host}:~$
- For root: root@{terminal_host}:~#
- Real permissions, real file sizes, real dates (Feb-Mar 2026), real PIDs, real kernel (6.1.0-18-amd64)
- No placeholders, no "...", no skipped output
- Each step's output must be realistic and complete

[CAPTION]
3-5 word caption for the experiment.
"""
        elif mode in ('general', 'language'):
            prompt = f"""You are an expert programming lab assistant. For this experiment aim:

"{aim}"

The user selected General Coding mode with this required code language: {target_language}

IMPORTANT GUIDELINES:
- Write the complete solution only in {target_language}.
- If the aim mentions any other programming language, ignore that language request and implement the same practical in {target_language}.
- Do not include code, syntax, headers, libraries, build tools, or examples from any other language.
- Keep comments minimal and only where genuinely needed.
- Provide a brief academic explanation of the core concepts being targeted.
- Show a realistic text output of running this {target_language} code.
- Do NOT include shell prompts like student@kali. Just show raw console execution outputs.

Respond EXACTLY in this format (use these exact tags):

[CONCEPT]
Write 3-4 lines explaining the concepts used. Academic style. Mention that the implementation uses {target_language}.

[CODE]
Write the full {target_language} source code. Plain text only, no markdown fences.

[OUTPUT]
Show REALISTIC output from running the code.
Make it look like a real terminal or console output. Do not show generic placeholder output.

[CAPTION]
Write a very short (3-5 words) descriptive caption for the output.
"""
        else:
            prompt = f"""You are an expert programming lab assistant. For this experiment aim:

"{aim}"

IMPORTANT GUIDELINES:
- Write clean and well-structured code.
- Provide a brief academic explanation of the core concepts being targeted.
- Show a realistic text output of running this code.
- If it's a programming language, provide the full source code.
- Do NOT include shell prompts like student@kali. Just show raw console execution outputs.

Respond EXACTLY in this format (use these exact tags):

[CONCEPT]
Write 3-4 lines explaining the concepts used. Academic style.

[CODE]
Write the code. Plain text only, no markdown fences.
Keep comments minimal — only where genuinely needed.

[OUTPUT]
Show REALISTIC output from running the code.
Make it look like a real terminal or console output. Do not show generic placeholder output.

[CAPTION]
Write a very short (3-5 words) descriptive caption for the output.
"""

        provider_attempts = [{
            'provider': provider,
            'model': model,
            'config': selected_provider_config,
            'keys': selected_api_keys,
        }]
        if provider == 'cerebras':
            groq_config = LLM_PROVIDERS['groq']
            provider_attempts.append({
                'provider': 'groq',
                'model': 'llama-3.3-70b-versatile',
                'config': groq_config,
                'keys': get_provider_keys(groq_config, ''),
            })
        if provider == 'freemodel_anthropic':
            fallback_model = selected_provider_config.get('fallback_model', 'claude-sonnet-4-6')
            if model != fallback_model:
                provider_attempts.append({
                    'provider': provider,
                    'model': fallback_model,
                    'config': selected_provider_config,
                    'keys': selected_api_keys,
                })

        # Retry with backoff for rate limiting (429)
        max_retries = 3
        for attempt in range(max_retries + 1):
            for provider_attempt in provider_attempts:
                attempt_provider = provider_attempt['provider']
                attempt_model = provider_attempt['model']
                attempt_config = provider_attempt['config']
                attempt_keys = provider_attempt['keys']

                for key_index, selected_key in enumerate(attempt_keys or ['']):
                    try:
                        text = create_chat_completion(
                            attempt_provider,
                            selected_key,
                            model=attempt_model,
                            messages=[{'role': 'user', 'content': prompt}],
                        )
                        break
                    except Exception as api_err:
                        err_str = str(api_err)
                        is_rate_limited = '429' in err_str or 'rate' in err_str.lower()
                        is_bad_key = (
                            '401' in err_str
                            or 'invalid api key' in err_str.lower()
                            or 'expired_api_key' in err_str.lower()
                            or 'authentication' in err_str.lower()
                        )
                        is_invalid_model = (
                            'invalid_request' in err_str.lower()
                            and 'model' in err_str.lower()
                        ) or '暂未开放' in err_str
                        has_backup_key = key_index < len(attempt_keys) - 1

                        if (is_rate_limited or is_bad_key) and has_backup_key:
                            print(f"{attempt_config['label']} key {key_index + 1} failed. Trying backup key...")
                            continue

                        if is_invalid_model and provider_attempt is not provider_attempts[-1]:
                            print(f"{attempt_config['label']} model {attempt_model} failed. Trying fallback model...")
                            break

                        if is_rate_limited and provider_attempt is not provider_attempts[-1]:
                            print(f"{attempt_config['label']} is busy. Falling back to {provider_attempts[-1]['config']['label']}...")
                            break

                        if is_rate_limited and attempt < max_retries:
                            wait_time = 15 * (2 ** attempt)  # 15s, 30s, 60s
                            print(f"Rate limited. Waiting {wait_time}s before retry {attempt + 1}/{max_retries}...")
                            time.sleep(wait_time)
                            break

                        raise

                if 'text' in locals():
                    break

            if 'text' in locals():
                break

        def extract_section(tag, text):
            pattern = rf"\[{tag}\](.*?)(?=\[(?:CONCEPT|CODE|PROCEDURE|OUTPUT|CAPTION)\]|$)"
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            
            if match:
                return match.group(1).strip()
            
            pattern_fallback = rf"(?:\*\*|##\s*)?{tag}(?:\*\*|:)?\s*\n(.*?)(?=\n(?:\*\*|##\s*)?(?:CONCEPT|CODE|PROCEDURE|OUTPUT|CAPTION)(?:\*\*|:)?\s*\n|\Z)"
            match_fallback = re.search(pattern_fallback, text, re.DOTALL | re.IGNORECASE)
            return match_fallback.group(1).strip() if match_fallback else None

        concept = extract_section("CONCEPT", text)
        caption = extract_section("CAPTION", text)
        if not concept: concept = "No concept description provided by API."
        if not caption: caption = "Experiment Output"

        result = {
            'concept': concept,
            'caption': caption,
            'mode': mode
        }
        if mode in ('general', 'language'):
            result['code_language'] = target_language

        if mode == 'os':
            procedure = extract_section("PROCEDURE", text) or "No procedure provided."
            procedure = re.sub(r'```[a-zA-Z]*', '', procedure).replace('```', '').strip()
            steps = parse_steps(procedure)
            result['steps'] = steps
            result['code'] = '\n\n'.join([f"Step {s['num']}: {s['explanation']}\n{s['command']}" for s in steps])
            result['output'] = '\n\n'.join([s['output'] for s in steps if s['output']])
        else:
            code = extract_section("CODE", text) or "// No code provided."
            output_part = extract_section("OUTPUT", text) or "No output provided."
            
            code = re.sub(r'```[a-zA-Z]*', '', code).replace('```', '').strip()
            output_part = re.sub(r'```', '', output_part).strip()

            result['code'] = code
            result['output'] = output_part

        return jsonify(result)
    except Exception as e:
        error_msg = str(e)
        status_code = 500
        if "401" in error_msg or "Invalid API Key" in error_msg or "Authentication" in error_msg:
            status_code = 401
        elif "403" in error_msg:
            status_code = 403
        elif "429" in error_msg or "Rate limit" in error_msg:
            status_code = 429
        return jsonify({'error': error_msg}), status_code



# ==================== API: Download .docx ====================
def set_font(paragraph, font_name='Times New Roman', size=12, bold=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        run.bold = bold


def add_bold_para(doc, text, font_name='Times New Roman', size=12, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = True
    set_font(p, font_name=font_name, size=size, bold=True)
    return p


def add_labeled_para(doc, label, content, font_name='Times New Roman', size=12):
    p = doc.add_paragraph()
    if any(c in content for c in ['*', '•', '·']) or '  ' in content:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = p.add_run(f'{label} ')
    run_label.bold = True
    run_label.font.name = font_name
    run_label.font.size = Pt(size)
    run_content = p.add_run(content)
    run_content.font.name = font_name
    run_content.font.size = Pt(size)
    return p


def add_code_para(doc, code_text, font_name='Times New Roman', size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_text)
    run.font.name = font_name
    run.font.size = Pt(size)


COMMAND_CAPTIONS = {
    'cal': 'Calendar Output',
    'cat': 'File Content Output',
    'cc': 'Compilation Output',
    'cd': 'Directory Change Output',
    'chmod': 'Permission Change Output',
    'chown': 'Ownership Change Output',
    'clear': 'Clear Screen Output',
    'cp': 'File Copy Output',
    'date': 'Date Command Output',
    'df': 'Disk Usage Output',
    'du': 'Storage Usage Output',
    'echo': 'Echo Command Output',
    'find': 'File Search Output',
    'free': 'Memory Usage Output',
    'gcc': 'Compilation Output',
    'grep': 'Search Command Output',
    'help': 'Shell Help Output',
    'history': 'Command History Output',
    'hostname': 'Hostname Output',
    'id': 'User Identity Output',
    'java': 'Program Execution Output',
    'javac': 'Compilation Output',
    'ls': 'Directory Listing Output',
    'make': 'Build Output',
    'man': 'Manual Page Output',
    'mkdir': 'Directory Creation Output',
    'mv': 'File Move Output',
    'ps': 'Process List Output',
    'pwd': 'Working Directory Output',
    'python': 'Program Execution Output',
    'python3': 'Program Execution Output',
    'reboot': 'Reboot Command Output',
    'rm': 'File Removal Output',
    'rmdir': 'Directory Removal Output',
    'sh': 'Script Execution Output',
    'stat': 'File Status Output',
    'su': 'User Switch Output',
    'sudo': 'Privilege Command Output',
    'top': 'Process Monitor Output',
    'touch': 'File Creation Output',
    'tty': 'Terminal Device Output',
    'uname': 'System Information Output',
    'uptime': 'System Uptime Output',
    'who': 'Logged-In Users Output',
    'whoami': 'Current User Output',
}


def make_caption_text(text, fallback='Terminal Output', max_words=8, max_chars=70):
    cleaned = re.sub(r'\s+', ' ', str(text or '')).strip()
    cleaned = re.sub(r'^Step\s*\d+\s*[:.)-]\s*', '', cleaned, flags=re.IGNORECASE)
    cleaned = cleaned.strip(' .:-')

    if not cleaned:
        return fallback

    if len(cleaned.split()) > max_words or len(cleaned) > max_chars:
        return fallback

    return cleaned or fallback


def make_step_caption(command, fallback='Terminal Output'):
    lines = [line.strip() for line in str(command or '').splitlines() if line.strip()]
    if not lines:
        return fallback

    first_line = re.sub(r'^[\w.-]+@[\w.-]+:.*?[#$]\s*', '', lines[0]).strip()
    first_line = re.sub(r'^\$\s*', '', first_line).strip()
    first_line = re.split(r'\s*(?:&&|\|\||;|\|)\s*', first_line, maxsplit=1)[0].strip()

    if not first_line:
        return fallback

    command_match = re.match(r'([./\w+-]+)', first_line)
    if not command_match:
        return fallback

    command_name = command_match.group(1).strip()
    base_name = command_name.replace('\\', '/').split('/')[-1].lower()

    if command_name.startswith('./'):
        return 'Program Execution Output'

    return COMMAND_CAPTIONS.get(base_name, f'Output of {base_name} Command')


def add_caption_para(doc, text, experiment_no, step_no=None, font_name='Times New Roman', size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_text = make_caption_text(text)
    if step_no:
        label = f'Figure {experiment_no}.{step_no} - {caption_text}'
    else:
        label = f'Figure {experiment_no} - {caption_text}'
    run = p.add_run(label)
    run.font.name = font_name
    run.font.size = Pt(size)


def get_terminal_lines(output_text, max_lines=28, max_line_chars=120):
    raw_lines = str(output_text or '').split('\n')
    lines = []
    shortened = len(raw_lines) > max_lines

    for raw_line in raw_lines[:max_lines]:
        line = raw_line.replace('\r', '')
        if len(line) > max_line_chars:
            line = line[:max_line_chars - 3] + '...'
            shortened = True
        lines.append(line)

    if shortened:
        lines.append(f'[Output shortened for export. Showing first {max_lines} lines.]')

    return lines or ['']


def create_terminal_image(output_text, img_width=600):
    # Image settings
    width = img_width
    font_size = 16
    padding = 20
    
    # Check for font
    try:
        font = ImageFont.truetype("consola.ttf", font_size) # Windows console font
    except IOError:
        try:
            font = ImageFont.truetype("cour.ttf", font_size) # Courier
        except IOError:
            font = ImageFont.load_default()

    # Calculate height
    lines = get_terminal_lines(output_text)
    line_height = font_size + 9 
    height = (len(lines) * line_height) + (2 * padding)
    
    # Create Image
    img = Image.new('RGB', (width, height), color=(0, 0, 0)) # Pure black background
    d = ImageDraw.Draw(img)
    
    # Draw text
    y = padding
    for line in lines:
        try:
            text_line = line
            # Normal font weight, matching Picture1.png color
            d.text((padding, y), text_line, font=font, fill=(201, 219, 213))
        except:
            pass
        y += line_height
        
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def add_normal_para(doc, text, font_name='Times New Roman', size=12, align=None):
    if align is None:
        if any(c in text for c in ['*', '•', '·']) or '  ' in text:
            align = WD_ALIGN_PARAGRAPH.LEFT
        else:
            align = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(p, font_name=font_name, size=size)
    return p


@app.route('/api/download', methods=['POST'])
def api_download():
    try:
        data = read_download_payload()
        if not data:
            return jsonify({'error': 'No data received for export.'}), 400
            
        experiments = data.get('experiments', [])
        if not experiments:
            return jsonify({'error': 'No experiment artifacts found to bundle.'}), 400
            
        settings = data.get('settings', {})

        font_name = settings.get('fontName', 'Times New Roman')
        body_size = int(settings.get('bodySize', 12))
        heading_size = int(settings.get('headingSize', 14))
        code_size = int(settings.get('codeSize', 10))
        caption_size = int(settings.get('captionSize', 10))
        image_width_inches = float(settings.get('imageWidth', 5.0))
        terminal_img_width = int(settings.get('terminalImgWidth', 600))
        output_filename = settings.get('outputFilename', 'Generated_Practical_File.docx')

        mode = data.get('mode', 'general')
        doc = Document()

        for i, exp in enumerate(experiments, 1):
            aim = exp.get('aim', 'N/A')
            concept = exp.get('concept', 'No concept description provided.')
            caption = exp.get('caption', 'Terminal Output Preview')
            steps = exp.get('steps', [])

            # Experiment heading
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'Experiment No. {i}')
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(heading_size)
            doc.add_paragraph('')

            add_labeled_para(doc, 'Aim:', aim, font_name, body_size)
            doc.add_paragraph('')
            
            if mode == 'os':
                add_bold_para(doc, 'Theory:', font_name, body_size)
                add_normal_para(doc, concept, font_name, body_size)
                doc.add_paragraph('')

                # Procedure with per-step output images
                add_bold_para(doc, 'Procedure:', font_name, body_size)

                if steps:
                    for step in steps:
                        step_num = step.get('num', '')
                        explanation = step.get('explanation', '')
                        command = step.get('command', '')
                        output = step.get('output', '')

                        # Step explanation
                        add_normal_para(doc, f"Step {step_num}: {explanation}", font_name, body_size)
                        
                        # For multi-line code (C programs etc.), show source code as text
                        is_multiline_code = command.count('\n') > 2
                        if is_multiline_code:
                            add_code_para(doc, command, font_name, code_size)

                        # Output as terminal image
                        if output.strip():
                            try:
                                img_buf = create_terminal_image(output, terminal_img_width)
                                pic_para = doc.add_paragraph()
                                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = pic_para.add_run()
                                run.add_picture(img_buf, width=Inches(image_width_inches))
                                add_caption_para(doc, make_step_caption(command), i, step_num, font_name, caption_size)
                            except Exception as img_err:
                                print(f"DEBUG: Step {step_num} image error: {img_err}")
                                add_code_para(doc, output, font_name, code_size)
                        
                        doc.add_paragraph('')  # spacing between steps
                else:
                    # Fallback
                    code = exp.get('code', '// No procedure available.')
                    output = exp.get('output', 'No output.')
                    add_code_para(doc, code, font_name, code_size)
                    doc.add_paragraph('')
                    add_bold_para(doc, 'Output:', font_name, body_size)
                    try:
                        img_buf = create_terminal_image(output, terminal_img_width)
                        pic_para = doc.add_paragraph()
                        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = pic_para.add_run()
                        run.add_picture(img_buf, width=Inches(image_width_inches))
                        add_caption_para(doc, caption, i, None, font_name, caption_size)
                    except Exception as img_err:
                        print(f"DEBUG: Error creating terminal image: {img_err}")
                        add_code_para(doc, output, font_name, code_size)

            else:
                code = exp.get('code', '// No code available.')
                output = exp.get('output', 'Program executed successfully.')

                add_labeled_para(doc, 'Concept Used:', concept, font_name, body_size)
                doc.add_paragraph('')
                add_bold_para(doc, 'Code:', font_name, body_size)
                add_code_para(doc, code, font_name, code_size)
                doc.add_paragraph('')
                add_bold_para(doc, 'Output:', font_name, body_size)

                try:
                    img_buf = create_terminal_image(output, terminal_img_width)
                    pic_para = doc.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = pic_para.add_run()
                    run.add_picture(img_buf, width=Inches(image_width_inches))
                    add_caption_para(doc, caption, i, None, font_name, caption_size)
                except Exception as img_err:
                    print(f"DEBUG: Error creating terminal image: {img_err}")
                    add_normal_para(doc, f'[Visual Output Unavailable - Log Trace follows]', font_name, body_size)
                    add_code_para(doc, output, font_name, code_size)

            if i < len(experiments):
                doc.add_page_break()

        file_buf = io.BytesIO()
        doc.save(file_buf)
        file_buf.seek(0)

        return send_file(
            file_buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_filename
        )
    except DownloadRequestError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        print(f"CRITICAL EXPORT ERROR: {traceback.format_exc()}")
        return jsonify({'error': f'Export Pipeline Fault: {str(e)}'}), 500


if __name__ == '__main__':
    print('\n  [+] PractiGen running at http://localhost:5000\n')
    app.run(host='0.0.0.0', port=5000, debug=True)
